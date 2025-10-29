from flask import Flask, request, session, redirect, url_for, render_template, jsonify, send_from_directory
from flask_mysqldb import MySQL
from flask import current_app
import MySQLdb.cursors
import openai
import pandas as pdf
import os
from docx import Document
import traceback
import requests 
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from io import BytesIO
from datetime import datetime
import re
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from dotenv import load_dotenv 
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement





app = Flask(__name__)
app.secret_key = 'b1dcf804bcada8cdf9300f13fc3e7d7e'  # Required for session management

app.config['MYSQL_HOST'] = 'mysql-2b9c8eba-chlakhna702-5683.h.aivencloud.com'  # Update this with your secret retrieval method
app.config['MYSQL_USER'] = 'avnadmin'  # Update this with your secret retrieval method
app.config['MYSQL_PASSWORD'] = 'AVNS_EMM0l9B433aCexcaKkt'  # Update this with your secret retrieval method
app.config['MYSQL_DB'] = 'dpa'  # Update this with your secret retrieval method
app.config['MYSQL_PORT'] = 10605 
mysql = MySQL(app)

load_dotenv()

# Set your OpenAI API key from environment variable
openai.api_key = os.getenv('OPENAI_API_KEY')

# OpenAI API Key
# openai.api_key = os.getenv("OPENAI_API_KEY")
@app.route('/')
@app.route('/login', methods = ['GET', 'POST'])

def login():
    message = ''
    if request.method == 'POST' and 'username' in request.form and 'password' in request.form:
        username = request.form['username']
        password = request.form['password']

        # Query to check for username and password in the database
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT * FROM users WHERE username = %s AND password = %s', (username, password))
        user = cursor.fetchone()

        if user:
            # If user exists, store information in the session
            session['loggedin'] = True
            session['userid'] = user['id']
            session['username'] = user['username']
            session['phone'] = user['phone']
            session['role'] = user['role']
            session['userlevel'] = user['userlevel']
            message = 'Logged in successfully!'
            return redirect(url_for('dashboard'))  # Redirect to the user profile or other page
            # return render_template('Dashboard.html', message=message)
        else:
            message = 'Incorrect username or password. Please try again.'

    return render_template('login.html', message=message)

@app.route('/logout')
def logout():
    # session.clear()  # Clear session
    # return redirect(url_for('logout'))  # Redirect to the login page
    return render_template('logout.html')

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    # Ensure user is logged in
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    # Get logged-in user's ID from the session
    user_id = session['userid']
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)

    # Fetch user information from the database
    cursor.execute('SELECT * FROM users WHERE id = %s', (user_id,))
    user = cursor.fetchone()

    if not user:
        return "User not found.", 404

    # If user submits the form to update their profile
    if request.method == 'POST':
        fullname = request.form['fullname']
        phone = request.form['phone']
        # role = request.form['role']
        # userlevel = request.form['userlevel']
        old_password = request.form['old_password_field']
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']

        # Validate old password
        if old_password and new_password:
            cursor.execute('SELECT password FROM users WHERE id = %s', (user_id,))
            db_password = cursor.fetchone()['password']
            if db_password != old_password:
                return "Old password is incorrect.", 400
            if new_password != confirm_password:
                return "New passwords do not match.", 400
            # Update the password
            cursor.execute('UPDATE users SET password = %s WHERE id = %s', (new_password, user_id))
        
        # Update other fields except username
        cursor.execute("""
            UPDATE users SET fullname = %s, phone = %s
            WHERE id = %s
        """, (fullname, phone, user_id))
        mysql.connection.commit()

        return redirect(url_for('register_list'))  # Redirect to the dashboard after updating

    cursor.close()
    return render_template('profile.html', user=user)




# Route for displaying users
@app.route('/register_list')
def register_list():
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    
    # Fetch all users from the database
    cursor.execute("SELECT * FROM users")
    users = cursor.fetchall()  # Fetch all the rows from the database
    cursor.close()

    # Pass the user data to the template
    return render_template('index.html', data=users)

@app.route('/form_one')
def form_one():
    # Your code for form_one here
    pass

# Route for registering a new user
@app.route('/register', methods=['GET', 'POST'])
def register():
    message = ''
    if request.method == 'POST':
        # Retrieve form data
        fullname = request.form['fullname']
        username = request.form['username']
        password = request.form['password']
        phone = request.form['phone']
        role = request.form['role']
        userlevel = request.form.get('userlevel', 'user')
        action = request.form.get('action')  # To know which button was clicked

        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)

        # Check if the username already exists
        cursor.execute('SELECT * FROM users WHERE username = %s', (username,))
        account = cursor.fetchone()

        if account:
            message = 'Account already exists!'
        elif not phone or not password or not username or not role:
            message = 'Please fill out all the required fields!'
        else:
            # Insert the new user into the database
            try:
                cursor.execute('INSERT INTO users (fullname, username, password, phone, role, userlevel) VALUES (%s ,%s, %s, %s, %s, %s)',
                               (fullname, username, password, phone, role, userlevel))
                mysql.connection.commit()
                message = 'You have successfully registered!'
            except MySQLdb.Error as e:
                print(f"Error occurred: {e}")
                message = 'Error occurred during registration. Please try again.'
            finally:
                cursor.close()

        # Check which button was clicked
        if action == 'create':
            # Redirect to index page after creating an account
            return redirect(url_for('register_list'))  # Redirect to the updated route
        elif action == 'add_account':
            # Stay on the form and show a success message
            return render_template('account.html', message=message)

    return render_template('account.html', message=message)


# Route for updating user information
@app.route('/update_account', methods=['POST'])
def update_account():
    if request.method == 'POST':
        # Retrieve form data for updating user information
        userid = request.form['userid']
        fullname = request.form['fullname']
        username = request.form['username']
        phone = request.form['phone']
        role = request.form['role']
        userlevel = request.form['userlevel']
        
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        try:
            cursor.execute("""
                UPDATE users 
                SET fullname=%s, username=%s,  phone=%s, role=%s, userlevel=%s 
                WHERE id=%s
            """, (fullname, username,  phone, role, userlevel, userid))
            mysql.connection.commit()
        except MySQLdb.Error as e:
            print(f"Error occurred: {e}")
            return jsonify({'error': 'Database error occurred'}), 500
        finally:
            cursor.close()

        return redirect(url_for('register_list'))  # Redirect to updated route


# Route for deleting a user
@app.route('/confirm_delete/<int:userid>', methods=['POST'])
def delete_account(userid):
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    try:
        cursor.execute("DELETE FROM users WHERE id=%s", (userid,))
        mysql.connection.commit()
    except MySQLdb.Error as e:
        print(f"Error occurred: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        cursor.close()
        
    return redirect(url_for('register_list'))  # Redirect back to updated route


# Route for fetching user data for editing
@app.route('/get_user/<int:userid>')
def get_user(userid):
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    try:
        cursor.execute("SELECT * FROM users WHERE id = %s", (userid,))
        user = cursor.fetchone()
        if user:
            return jsonify({
                'id': user['id'],
                'fullname':user['fullname'],
                'username': user['username'],
                'phone': user['phone'],
                'role': user['role'],
                'userlevel': user['userlevel']
            })
        else:
            return jsonify({'error': 'User not found'}), 404
    except MySQLdb.Error as e:
        print(f"Error occurred: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        cursor.close()

@app.route('/edit', methods=['GET', 'POST'])
def edit():
    message = ''

    if request.method == 'POST':
        # Get form data from the POST request
        new_token = request.form.get('new_token')
        new_groupid = request.form.get('new_groupid')

        # Only proceed if at least one value is provided
        if new_token or new_groupid:
            cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
            
            # Prepare the update query
            update_query = "UPDATE telegbot SET "
            params = []
            
            if new_token:
                update_query += "token = %s, "
                params.append(new_token)
            
            if new_groupid:
                update_query += "groupid = %s, "
                params.append(new_groupid)
            
            update_query = update_query.rstrip(', ')  # Trim the trailing comma
            update_query += " LIMIT 1"  # Assuming you want to update the first record
            
            cursor.execute(update_query, tuple(params))
            mysql.connection.commit()
            cursor.close()

            message = 'Information updated successfully!'
            return redirect(url_for('edit'))

    # Fetch the existing token and groupid from the database to prepopulate the form
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cursor.execute("SELECT token, groupid FROM telegbot LIMIT 1")
    telegbot_data = cursor.fetchone()
    cursor.close()

    # Pass the data to the template
    return render_template('settings.html', message=message, telegbot_data=telegbot_data)





@app.route('/formtwo')
def formtwo():
    # Correctly formatted Google Sheets export URL for CSV
    csv_url ='https://docs.google.com/spreadsheets/d/e/2PACX-1vS-HfWE1g10jIyJWjZjGRTfGZ_kQGqphrl9aHBSs_1Q_oXgT6leUu-Nv9wE64R4Pk-Pw4vdUr52gQD2/pub?output=csv'
    


    try:
        # Fetch the CSV data with robust error handling
        df = pd.read_csv(
            csv_url,
            sep=',',  # CSV delimiter
            engine='python',  # Use Python engine for flexible parsing
            on_bad_lines='skip'  # Skip problematic rows (requires Pandas >= 1.3.0)
        )

        # Extract column names and data
        column_names = df.columns.tolist()
        data = df.fillna('').values.tolist()  # Replace NaN with empty strings

        # Pass the data and column names to the HTML template
        return render_template('formtwo.html', data=data, column_names=column_names)

    except Exception as e:
        # Return error message for debugging
        return f"Error fetching data: {e}"



@app.route('/dashboard')
def dashboard():
    if 'loggedin' in session:
        return render_template('Dashboard.html')
    return redirect(url_for('login')) 


def generate_report_with_chatgpt(data, user_prompt):
    try:
        # Token limits for the model
        total_token_limit = 200000  # Maximum input + output tokens
        max_output_tokens = 16000  # Reserve 16,000 tokens for output
        input_token_limit = total_token_limit - max_output_tokens  # Input token space
        chunk_size = input_token_limit - 1000  # Keep chunk size smaller to account for prompt length

        # Function to split data into manageable chunks
        def split_data_into_chunks(data, chunk_size):
            return [data[i:i + chunk_size] for i in range(0, len(data), chunk_size)]

        # Split the data into chunks
        chunks = split_data_into_chunks(data, chunk_size)

        # Process only the last chunk
        last_chunk = chunks[-1]
        prompt = f""" 
        You are a data analyst. Analyze the following dataset and provide insights based on the task provided:

        Dataset:
        {last_chunk}

        Task:
        {user_prompt}
        """

        # Ensure prompt fits within input token limit
        if len(prompt) > input_token_limit:
            raise ValueError(f"Prompt for the last chunk exceeds the input token limit of {input_token_limit} tokens.")

        # Make the API call to GPT for the last chunk
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a data analyst."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=max_output_tokens,
            temperature=0.7
        )

        # Extract and return the generated content for the last chunk
        last_chunk_result = response['choices'][0]['message']['content'].strip()
        return last_chunk_result

    except Exception as e:
        error_message = f"Error: {str(e)}\n{traceback.format_exc()}"
        print(error_message)
        return error_message


@app.route('/form_two')
def form_two():
    # Correctly formatted Google Sheets export URL for CSV
    csv_url ='https://docs.google.com/spreadsheets/d/e/2PACX-1vS-HfWE1g10jIyJWjZjGRTfGZ_kQGqphrl9aHBSs_1Q_oXgT6leUu-Nv9wE64R4Pk-Pw4vdUr52gQD2/pub?output=csv'

    try:
        # Fetch the CSV data with robust error handling
        df = pd.read_csv(
            csv_url,
            sep=',',  # CSV delimiter
            engine='python',  # Use Python engine for flexible parsing
            on_bad_lines='skip'  # Skip problematic rows (requires Pandas >= 1.3.0)
        )

        # Extract column names and data
        column_names = df.columns.tolist()
        data = df.fillna('').values.tolist()  # Replace NaN with empty strings

        # Pass the data and column names to the HTML template
        return render_template('report2.html', data=data, column_names=column_names)

    except Exception as e:
        # Return error message for debugging
        return f"Error fetching data: {e}"


@app.route('/generate-report', methods=['POST'])
def generate_report():
    try:
        # Ensure the request is JSON
        if not request.is_json:
            return jsonify({"error": "Request must be JSON"}), 400

        # Get the user prompt from the POST request JSON body
        request_data = request.get_json()
        user_prompt = request_data.get('user_prompt')

        if not user_prompt:
            return jsonify({"error": "User prompt is missing"}), 400

        # Correctly formatted Google Sheets export URL for CSV
        csv_url ='https://docs.google.com/spreadsheets/d/e/2PACX-1vS-HfWE1g10jIyJWjZjGRTfGZ_kQGqphrl9aHBSs_1Q_oXgT6leUu-Nv9wE64R4Pk-Pw4vdUr52gQD2/pub?output=csv'
        # Fetch the CSV data from Google Sheets
        try:
            df = pd.read_csv(
                csv_url,
                sep=',',  # CSV delimiter
                engine='python',  # Use Python engine for flexible parsing
                on_bad_lines='skip'  # Skip problematic rows (requires Pandas >= 1.3.0)
            )
        except Exception as e:
            return jsonify({"error": f"Error fetching data from Google Sheets: {e}"}), 500

        # Convert the DataFrame to a string format suitable for GPT
        data_as_text = df.to_markdown(index=False)  # Convert to Markdown table
        full_prompt = f"The following is the dataset:\n\n{data_as_text}\n\n{user_prompt}"

        # Generate the report using GPT-4
        generated_report_content = generate_report_with_chatgpt(data_as_text, user_prompt)

        if generated_report_content:
            # Save the report as a Word document
            word_filename = 'generated_report.docx'
            save_report_as_word(generated_report_content, word_filename)

            # Format the report content as HTML for display
            formatted_report = format_report_as_html(generated_report_content)

            # Return both the HTML-formatted report and the filename for the downloadable Word document
            return jsonify({"report": formatted_report, "filename": word_filename})
        else:
            return jsonify({"error": "Report generation failed"}), 500

    except Exception as e:
        return jsonify({"error": str(e)}), 500
# def generate_report():
#     try:
#         # Ensure the request is JSON
#         if not request.is_json:
#             return jsonify({"error": "Request must be JSON"}), 400

#         # Get the user prompt from the POST request JSON body
#         request_data = request.get_json()
#         user_prompt = request_data.get('user_prompt')

#         if not user_prompt:
#             return jsonify({"error": "User prompt is missing"}), 400

#         # Correctly formatted Google Sheets export URL for CSV
#         csv_url ='https://docs.google.com/spreadsheets/d/e/2PACX-1vS-HfWE1g10jIyJWjZjGRTfGZ_kQGqphrl9aHBSs_1Q_oXgT6leUu-Nv9wE64R4Pk-Pw4vdUr52gQD2/pub?output=csv'
#         # Fetch the CSV data from Google Sheets
#         try:
#             df = pd.read_csv(
#                 csv_url,
#                 sep=',',  # CSV delimiter
#                 engine='python',  # Use Python engine for flexible parsing
#                 on_bad_lines='skip'  # Skip problematic rows (requires Pandas >= 1.3.0)
#             )
         
#         except Exception as e:
#             return jsonify({"error": f"Error fetching data from Google Sheets: {e}"}), 500

#         # Convert the DataFrame to a string format suitable for GPT
#         data_as_text = df.to_markdown(index=False)  # Convert to Markdown table
#         full_prompt = f"The following is the dataset:\n\n{data_as_text}\n\n{user_prompt}"

#         # Generate the report using GPT-4
#         generated_report_content = generate_report_with_chatgpt(data_as_text, user_prompt)

#         if generated_report_content:
#             # Save the report as a Word document
#             word_filename = 'generated_report.docx'
#             save_report_as_word(generated_report_content, word_filename)

#             # Format the report content as HTML for display
#             formatted_report = format_report_as_html(generated_report_content)

#             # Return both the HTML-formatted report and the filename for the downloadable Word document
#             return jsonify({"report": formatted_report, "filename": word_filename})
#         else:
#             return jsonify({"error": "Report generation failed"}), 500

#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

def format_report_as_html(report_content):
    formatted_report = ""
    in_list = False
    in_numbered_list = False
    in_table = False
    in_code_block = False  # Track whether we're inside a code block
    row_counter = 0  # Counter to track rows

    # Iterate over each line of the report content
    for line in report_content.split('\n'):
        line = line.rstrip()  # Retain trailing spaces in code blocks

        # Handle code block opening and closing (triple backticks or ``` for code block)
        if line.startswith("```"):
            if not in_code_block:
                # Opening a code block with a background color
                formatted_report += "<pre style='background-color: #000000; width: 100%; padding: 10px; border-radius: 8px;'><code class='language-java' style='color: #FFFFFF;'>"
                in_code_block = True
            else:
                # Closing a code block
                formatted_report += "</code></pre>"
                in_code_block = False
            continue

        # If we're in a code block, preserve indentation and line breaks
        if in_code_block:
            formatted_report += line + "\n"
            continue

        # Skip separator lines (|----|)
        if re.match(r'^\|[-\s]+\|$', line):  # Matches separator lines like |----|
            continue  # Ignore lines that are table separators (lines with just dashes)

        # Detect table headers (starting with '|') and table rows
        if line.startswith("|"):  # Process table rows or headers
            row_counter += 1  # Increment the row counter

            # Skip the second row (assumed separator row)
            if row_counter == 2:
                continue

            cells = [cell.strip() for cell in line.split('|') if cell]  # Split the line into cells

            if not in_table:
                formatted_report += "<table border='1' style='width:100%; text-align:left; border-collapse: collapse; background-color: #f1f8e9;'>"
                in_table = True

            # If it's the first row (header), format it as the table header with a background color
            if row_counter == 1:
                formatted_report += "<thead style='background-color: #388e3c; color: black;'><tr>" + ''.join([f"<th style='padding: 10px;'>{cell}</th>" for cell in cells]) + "</tr></thead><tbody>"
            else:
                formatted_report += "<tr>" + ''.join([f"<td style='padding: 10px; border: 1px solid black;'>{cell}</td>" for cell in cells]) + "</tr>"

        # Close the table after encountering an empty line
        elif in_table and not line:
            formatted_report += "</tbody></table>"
            in_table = False

        # Handle different heading levels (H1, H2, H3, etc.) with vibrant colors
        elif line.startswith("# "):
            formatted_report += f"<h1 style='font-size: 2em; text-align: left; font-weight: bold; color: black; margin-top: 24px;'>{line[2:]}</h1>"
            in_list = False
            in_numbered_list = False

        elif line.startswith("## "):
            formatted_report += f"<h2 style='font-size: 1.75em; text-align: left; font-weight: bold; color: black; margin-top: 22px;'>{line[3:]}</h2>"
            in_list = False
            in_numbered_list = False

        elif line.startswith("### "):
            formatted_report += f"<h3 style='font-size: 1.5em; text-align: left; font-weight: bold; color: black; margin-top: 20px;'>{line[4:]}</h3>"
            in_list = False
            in_numbered_list = False

        elif line.startswith("#### "):
            formatted_report += f"<h4 style='font-size: 1.25em; text-align: left; font-weight: bold; color: black; margin-top: 18px;'>{line[5:]}</h4>"
            in_list = False
            in_numbered_list = False

        # Bullet points (left aligned, indented)
        elif line.startswith("- "):
            if not in_list:
                formatted_report += "<ul style='text-align: left; padding-left: 24px; color: black;'>"
                in_list = True
            formatted_report += f"<li>{apply_bold_and_italic(line[2:])}</li>"

        # Numbered lists (left aligned, indented)
        elif re.match(r'^\d+\.\s', line):
            if not in_numbered_list:
                formatted_report += "<ol style='text-align: left; padding-left: 24px; color: black;'>"
                in_numbered_list = True
            formatted_report += f"<li>{apply_bold_and_italic(line[3:])}</li>"

        # General paragraph (aligned left and applying bold/italic formatting)
        else:
            if in_list:
                formatted_report += "</ul>"
                in_list = False
            if in_numbered_list:
                formatted_report += "</ol>"
                in_numbered_list = False
            formatted_report += f"<p style='text-align: left; color: black;'>{apply_bold_and_italic(line)}</p>"

    # Ensure that any open lists, tables, or code blocks are closed
    if in_list:
        formatted_report += "</ul>"
    if in_numbered_list:
        formatted_report += "</ol>"
    if in_table:
        formatted_report += "</tbody></table>"
    if in_code_block:
        formatted_report += "</code></pre>"

    return formatted_report

# Function to handle bold and italic formatting
def apply_bold_and_italic(line):
    # Convert **text** to <b>text</b> and *text* to <i>text</i>
    line = re.sub(r'\*\*(.+?)\*\*', r'<b style="color: black;">\1</b>', line)  # Bold text in a vibrant pink
    line = re.sub(r'\*(.+?)\*', r'<i style="color: black;">\1</i>', line)  # Italic text in a blue
    return line
    

def apply_bold(run):
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # Vibrant pink for bold

def apply_italic(run):
    run.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # Blue for italic

# Function to handle markdown-like text formatting in Word
def add_markdown_formatted_text(paragraph, text):
    bold_italic_pattern = re.compile(r'\*\*\*(.+?)\*\*\*')
    bold_pattern = re.compile(r'\*\*(.+?)\*\*')
    italic_pattern = re.compile(r'\*(.+?)\*')

    cursor = 0
    while cursor < len(text):
        nearest_match = None
        nearest_style = None
        nearest_start = len(text)

        for style, pattern in [("bold_italic", bold_italic_pattern), ("bold", bold_pattern), ("italic", italic_pattern)]:
            match = pattern.search(text, cursor)
            if match and match.start() < nearest_start:
                nearest_match = match
                nearest_style = style
                nearest_start = match.start()

        if nearest_match:
            if nearest_start > cursor:
                paragraph.add_run(text[cursor:nearest_start])
            if nearest_style == "bold_italic":
                run = paragraph.add_run(nearest_match.group(1))
                apply_bold(run)
                apply_italic(run)
            elif nearest_style == "bold":
                apply_bold(paragraph.add_run(nearest_match.group(1)))
            elif nearest_style == "italic":
                apply_italic(paragraph.add_run(nearest_match.group(1)))
            cursor = nearest_match.end()
        else:
            paragraph.add_run(text[cursor:])
            break

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Function to add styled headings to Word
def add_styled_heading(doc, text, level, color):
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = color

# Function to create a styled table
def create_styled_table(doc, data):
    table = doc.add_table(rows=len(data), cols=len(data[0]), style="Table Grid")
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx].cells
        for cell_idx, cell_data in enumerate(row_data):
            row_cells[cell_idx].text = cell_data
            row_cells[cell_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
from docx.oxml import OxmlElement
from docx.oxml import OxmlElement

def save_report_as_word(report_content, filename):
    try:
        doc = Document()

        create_cover_page(doc, "Generated Report")
        
        # Adjust margins
        for section in doc.sections:
            section.top_margin = Pt(56.7)
            section.bottom_margin = Pt(56.7)
            section.left_margin = Pt(56.7)
            section.right_margin = Pt(56.7)

        lines = report_content.split('\n')
        table = None
        row_counter = 0  # Counter to track which row is being processed

        for line in lines:
            line = line.strip()

            # Add page breaks for "---"
            if line == "---":
                doc.add_page_break()
            # Headings
            elif line.startswith("# "):
                heading = doc.add_heading(line[2:], level=1)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # H1 in orange-red
            elif line.startswith("## "):
                heading = doc.add_heading(line[3:], level=2)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # H2 in blue
            elif line.startswith("### "):
                heading = doc.add_heading(line[4:], level=3)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # H3 in green
            elif line.startswith("#### "):
                heading = doc.add_heading(line[5:], level=4)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # H4 in purple
            # Bullet points
            elif line.startswith("* "):
                bullet_paragraph = doc.add_paragraph(line[2:], style="List Bullet")
                bullet_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Bold and italic formatting
            elif "**" in line or "*" in line:
                paragraph = doc.add_paragraph()
                add_markdown_formatted_text(paragraph, line)
            # Tables
            elif "|" in line:
                table_data = [cell.strip() for cell in line.split("|") if cell]
                row_counter += 1  # Increment row counter for each table line

                # Skip the second row (separator row)
                if row_counter == 2:
                    continue

                if not table:
                    # Create table and add header row
                    table = doc.add_table(rows=1, cols=len(table_data), style="Table Grid")
                    hdr_cells = table.rows[0].cells
                    for i, cell_data in enumerate(table_data):
                        hdr_cells[i].text = cell_data
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # White header text
                        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Apply background color for header
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), "FFFFFF")  # Blue background for header
                        hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                else:
                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(table_data):
                        row_cells[i].text = cell_data
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # General text
            else:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.font.color.rgb = RGBColor(66, 66, 66)  # Dark grey text
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Save the document
        save_path = os.path.join(os.getcwd(), 'static/reports', filename)
        if not os.path.exists(os.path.dirname(save_path)):
            os.makedirs(os.path.dirname(save_path))

        doc.save(save_path)
        print(f"File '{filename}' saved successfully.")

    except Exception as e:
        print(f"Error saving report: {e}")


def create_cover_page(doc, report_title):
    section = doc.sections[0]
    section.page_height = Pt(850)  # A4 size height
    section.page_width = Pt(595)   # A4 size width

    section = doc.sections[0]
    header = section.header

    # Add a new paragraph for the text in the header
    text_paragraph = header.add_paragraph()

    # Add "Draft Report" on the left
    run_left = text_paragraph.add_run("Draft Report")
    run_left.font.size = Pt(12)
    run_left.font.bold = True
    run_left.font.color.rgb = RGBColor(0, 70, 132)

    # Add a tab for right alignment
    text_paragraph.add_run("\t")

    # Add "Baseline Survey" on the right
    run_right = text_paragraph.add_run("Baseline Survey")
    run_right.font.size = Pt(12)
    run_right.font.bold = True
    run_right.font.color.rgb = RGBColor(0, 70, 132)

    # Set up a tab stop for right alignment
    tab_stops = text_paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6), alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)  # Adjust `Inches(6)` to match your page width

    # Adjust the paragraph alignment to justify so tab stops are respected
    text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
 
    # Add another paragraph for the horizontal line
    line_paragraph = header.add_paragraph()  # Add a new paragraph for the line
    line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the line

    # Add the line
    run_line = line_paragraph.add_run("_" * 50)  # Creates the horizontal line
    run_line.font.size = Pt(25)  # Adjust line thickness by changing font size
    run_line.font.color.rgb = RGBColor(0, 70, 132)
    

    # Set margins to 2 cm (56.7 points)
    for section in doc.sections:
        section.top_margin = Pt(56.7)
        section.bottom_margin = Pt(56.7)
        section.left_margin = Pt(56.7)
        section.right_margin = Pt(56.7)

    # Add logo from URL
    
    
    # base_dir = r"d:\Clone\test\Fisheries-Information-Management-System-FIMS-"
    base_dir = os.path.abspath(os.path.dirname(__file__))

    image_paths = [
    os.path.join(base_dir, "static", "img", "profile_images", "goverment2.jpg"),
    os.path.join(base_dir, "static", "img", "profile_images", "2.jpg"),
    os.path.join(base_dir, "static", "img", "profile_images", "adic.jpg"),
    os.path.join(base_dir, "static", "img", "profile_images", "dcx.jpg"),
    ]

    # Verify all images exist
    for image_path in image_paths:
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"File not found: {image_path}")


    # Add all 4 images in a single paragraph (side by side)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align all images in the center

    for image_path in image_paths:
        with open(image_path, "rb") as img_file:
            image_data = BytesIO(img_file.read())
            run = paragraph.add_run()
            picture = run.add_picture(image_data, width=Pt(120), height=Pt(50))

            picture_width = Pt(120)  # Set desired width
            picture_height = Pt(50)


    base_dir = os.path.abspath(os.path.dirname(__file__))
    logo_path = os.path.join(base_dir, "static", "img", "profile_images", "people.jpg")

    if os.path.exists(logo_path):
    # Open the image file directly
        with open(logo_path, "rb") as img_file:
            logo_image = BytesIO(img_file.read())

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        picture = run.add_picture(logo_image, width=Pt(480))  # Adjust width as needed

        picture_width = Pt(480)  # Set desired width
        picture_height = Pt(200)

        


    # Add some spacing
    # doc.add_paragraph("\n")
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = cover.add_run("The Land Allocation for Social and Economic Development ")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.bold = True
    run.add_break()
    run = cover.add_run("Project III (LASED III) and the Cambodia Sustainable")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.bold = True
    run.add_break()
    run = cover.add_run("Livelihood for Indigenous Communities Project (CSLICP)")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.bold = True
    # cover.add_run("\n").font.size = Pt(24)

    # for _ in range(2):
    #     doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    # Add prepared for
    prepared_for = doc.add_paragraph()
    prepared_for.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prepared_for.add_run("Draft Report – Baseline Survey")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.bold = True

    # Place the date at the bottom center of the cover page
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run(f"{datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.bold = True

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.footer.is_linked_to_previous = False
    new_section.header.is_linked_to_previous = False


@app.route('/download-report/<filename>', methods=['GET'])
def download_report(filename):
    try:
        directory = os.path.join(os.getcwd(), 'static/reports')
        return send_from_directory(directory=directory, path=filename, as_attachment=True)
    except Exception as e:
        return str(e)



@app.route('/send-to-telegram', methods=['POST'])
def send_to_telegram():
    try:
        # Create a cursor object to retrieve the token and groupid from telegbot table
        cursor = mysql.connection.cursor()
        cursor.execute("SELECT token, groupid FROM telegbot WHERE id = 1")  # Fetch first row

        result = cursor.fetchone()

        # Close the cursor
        cursor.close()

        if result is None:
            print("No token or group ID found in the database.")
            return jsonify({"error": "No token or group ID found in the database."}), 404

        # Extract token and groupid from the result
        token = result[0]
        group_id = result[1]
        print(f"Token: {token}, Group ID: {group_id}")

        # Path to the report file
        file_path = os.path.join(os.getcwd(), 'static/reports/generated_report.docx')

        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return jsonify({"error": "Report file not found."}), 404

        # Telegram API URL for sending messages
        telegram_api_url = f"https://api.telegram.org/bot{token}/sendDocument"

        # Send the file to the Telegram group using requests
        with open(file_path, 'rb') as document:
            response = requests.post(telegram_api_url, data={'chat_id': group_id}, files={'document': document})
            print(f"Telegram API Response: {response.status_code} - {response.text}")

        if response.status_code == 200:
            return jsonify({"success": "Report sent to Telegram successfully!"})
        else:
            error_message = f"Telegram API error: {response.status_code} - {response.text}"
            print(error_message)
            return jsonify({"error": error_message}), 500

    except Exception as e:
        error_message = f"Exception occurred: {str(e)}"
        print(error_message)
        return jsonify({"error": error_message}), 500


@app.route('/report')
def report():
    return render_template('Report.html')





if __name__ == "__main__":
    app.run()
